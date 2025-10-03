import os
import json
from datetime import datetime
from tempfile import NamedTemporaryFile
from flask import (
    Flask, render_template, request, jsonify, session,
    redirect, url_for, send_file, abort
)
import google.generativeai as genai
from docx import Document
from werkzeug.utils import secure_filename
import yaml
from flask_session import Session
from bs4 import BeautifulSoup
import markdown



app = Flask(__name__)
app.config['SESSION_TYPE'] = 'filesystem'   # store sessions in /flask_session folder
app.config['SESSION_FILE_DIR'] = os.path.join(app.root_path, 'flask_session')
os.makedirs(app.config['SESSION_FILE_DIR'], exist_ok=True)
app.config['SESSION_PERMANENT'] = False
Session(app)
app.secret_key = "my-very-strong-hardcoded-secret"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/ai-generate-draft')
def ai_generate_draft():
    return render_template('aidraft.html')


# ---------- NEW: Config ----------
# Expect: export GEMINI_API_KEY="..."
with open("keys.yaml", "r") as f:
    keys = yaml.safe_load(f)
GEMINI_API_KEY = keys.get("GEMINI_API_KEY")
GEMINI_MODEL = "models/gemini-2.5-pro"

# Root/contracts/contract_template.docx  (adjust name if needed)
TEMPLATE_PATH = r"C:\My Folder\projects\legalbook-new-clm\contracts\General_Contract_Template.docx"

SYSTEM_INSTRUCTION = """You are a senior contracts counsel.
Given (1) structured deal terms and (2) a base contract template, draft a complete, legally sound contract.
- Use the template’s structure/sections and fill in placeholders.
- Insert the provided parties: our company, counterparty, and any additional_parties if present.
    - For each additional_party, treat its "type" field as the role or identity heading (e.g., Witness, Guarantor, Arbitrator, Affiliate). 
    - If no additional_parties are provided, ignore this step.
    - Always reflect their name, email, and registered office address faithfully in the contract.
- Keep numbering consistent and professional.
- If you identify missing protections, compliance items, or industry-standard clauses, you may add them proactively (e.g., indemnity, liability caps, data protection, termination rights, governing law, dispute resolution).
- Analyze our requirements, party details, and documents to suggest optimal clauses, identify potential risks, and ensure compliance with industry standards.
- Do not remove user-provided information, but you may enhance or expand the contract where appropriate.
- Keep the tone professional, concise, and enforceable.
- Output clean, editable text with clear headings and sections (no markdown fences).
"""


# Initialize Gemini model lazily (only when key exists)
def get_model():
    if not GEMINI_API_KEY:
        raise RuntimeError("Missing GEMINI_API_KEY")
    genai.configure(api_key=GEMINI_API_KEY)
    return genai.GenerativeModel(GEMINI_MODEL, system_instruction=SYSTEM_INSTRUCTION)

def read_template_text(path: str) -> str:
    """Read a .docx template and return plain text (paragraphs)."""
    if not os.path.exists(path):
        return ""
    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        return "\n".join(parts).strip()
    except Exception:
        return ""

# ---------- NEW: Generate Contract (POST from Step 4 button) ----------
@app.route('/generate-contract', methods=['POST'])
def generate_contract():
    try:
        data = request.get_json(force=True)
    except Exception:
        return "Invalid JSON", 400
    # Basic sanity
    title = (data.get("contract_title") or "Contract Draft").strip()
    session['generated_contract_title'] = title

    # 1) Read template text (optional but recommended)
    template_text = read_template_text(TEMPLATE_PATH)

    # 2) Build prompt content
    # Send all fields as JSON so the model can fill them deterministically
    deal_terms_json = json.dumps(data, indent=2)

    prompt = f"""
BASE TEMPLATE (DOCX -> text):
{template_text if template_text else "[No base template file found or couldn't be read]"}

DEAL TERMS (JSON, includes additional_parties if provided):
{deal_terms_json}


TASK:
Combine the BASE TEMPLATE and the DEAL TERMS into a complete contract, filling in all specific details.
- Replace placeholders like [Party A], [Party B], dates, values, or payment terms with the DEAL TERMS.
- If "additional_parties" are provided in the JSON, include them explicitly in the contract. 
  Use their "type" field as the legal identity or role heading (e.g., Witness, Guarantor, Arbitrator, Affiliate). 
  List their name, email, and registered office appropriately. 
- If "additional_parties" is empty or missing, omit this step.
- Ensure clauses that reflect 'special_clauses' are included.
- Provide a final, clean, ready-to-edit contract text.
"""

    # 3) Call Gemini
    try:
        model = get_model()
        resp = model.generate_content(prompt)
        text = (resp.text or "").strip()
        if not text:
            return "LLM returned empty response", 502
    except Exception as e:
        return f"LLM error: {e}", 502

    html_content = markdown.markdown(
        text,
        extensions=["extra", "nl2br", "sane_lists"]
    )
    # 4) Stash in session and tell client where to go
    session['generated_contract_html'] = html_content
    return jsonify({"redirect_url": url_for('edit_contract')}), 200


# ---------- NEW: Editor page to review & edit ----------
@app.route('/edit-contract')
def edit_contract():
    content = session.get('generated_contract_html')
    if not content:
        # If nothing generated yet, send back to AI Draft
        return redirect(url_for('ai_generate_draft'))
    title = session.get('generated_contract_title', 'Contract Draft')
    return render_template('edit_contract.html', content=content, title=title)


# ---------- NEW: Download as .docx ----------
@app.route('/download-contract', methods=['POST'])
def download_contract():
    content = (request.form.get('content') or "").strip()
    if not content:
        return "No content to download", 400

    title = session.get('generated_contract_title', 'Contract Draft')
    safe_name = secure_filename(title) or "Contract_Draft"
    filename = f"{safe_name}.docx"

    # Build a basic .docx from the plain text
    doc = Document()
    soup = BeautifulSoup(content, "html.parser")

    for block in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        if block.name in ["h1", "h2", "h3"]:
            level = int(block.name[1])
            doc.add_heading(block.get_text(strip=True), level=level)

        elif block.name == "p":
            para = doc.add_paragraph()
            for child in block.children:
                if child.name == "strong":
                    run = para.add_run(child.get_text())
                    run.bold = True
                elif child.name == "em":
                    run = para.add_run(child.get_text())
                    run.italic = True
                elif child.name == "u":
                    run = para.add_run(child.get_text())
                    run.underline = True
                elif child.string:
                    para.add_run(child.string)

        elif block.name == "li":
            para = doc.add_paragraph("• " + block.get_text(strip=True))


    # Save to a temp file and return
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    return send_file(
        tmp_path,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



if __name__ == '__main__':
    app.run(debug=True)